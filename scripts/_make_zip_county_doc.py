"""One-shot generator for docs/zip-county-matching.docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
OUT  = ROOT / "docs" / "zip-county-matching.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def H(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def P(text, *, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def CODE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    return p


def BULLET(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


# ── Title ─────────────────────────────────────────────────────────────────
title = doc.add_heading("ZIP ↔ County Matching Pipeline", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

P("How foreclosure / Realtor / map data are tied together via ZIP and county FIPS.",
  italic=True)

# ── Overview ──────────────────────────────────────────────────────────────
H("Overview", 1)
P(
    "The pipeline builds a single ZIP→FIPS crosswalk from raw geospatial sources, "
    "lifts it into Postgres as a permanent lookup table, and uses it to enrich "
    "foreclosure rows so they can be joined to Realtor county/ZIP metrics. The "
    "matching is geometric (point-in-polygon), not name-based, so it survives "
    "name spelling differences and county boundary edge cases."
)

# ── Data sources ──────────────────────────────────────────────────────────
H("Data sources", 1)

H("1. County polygons — counties-fips.geojson", 2)
BULLET("Path: data/counties-fips.geojson")
BULLET("One Feature per US county, tagged with 5-digit GEOID/FIPS.")
BULLET("Origin: US Census TIGER counties, packaged as GeoJSON.")
BULLET("Role: master shapefile for point-in-polygon county assignment.")

H("2. ZCTA polygons — OpenDataDE (per state)", 2)
BULLET("Source: https://github.com/OpenDataDE/State-zip-code-GeoJSON")
BULLET("URL pattern: raw.githubusercontent.com/OpenDataDE/State-zip-code-GeoJSON/master/<abbr>_<state>_zip_codes_geo.min.json")
BULLET("Cached locally at: data/_zcta_cache/<abbr>_<state>.geojson")
BULLET("ZIP field discovered as ZCTA5CE20 / ZCTA5CE10 / ZCTA5 / GEOID / ZIP / ZIPCODE.")
BULLET("Role: provides the polygon for each ZIP so its centroid can be located.")

H("3. Realtor county metrics — realtor_inventory_county (Neon)", 2)
BULLET("Loaded by scripts/import_realtor_county_state.py from "
       "Mailer_Data_set/Realtor/Monthly Housing Inventory/Monthly_Inventory__County.csv.")
BULLET("Key: (month_date_yyyymm, county_fips). FIPS is zero-padded to 5 chars on import.")
BULLET("Role: source of human-readable county_name used to fill zip_county_map.county_name.")

H("4. Foreclosure inventory — foreclosure_records (Neon)", 2)
BULLET("~882k rows. source ∈ { 'foreclosure_com', 'auction_com' }.")
BULLET("zip is always populated for foreclosure_com; county is set only for auction_com (as a name, not a FIPS).")
BULLET("Role: the table the crosswalk enriches — a county name is backfilled where missing.")

# ── Stage 1: build crosswalk ─────────────────────────────────────────────
H("Stage 1 — Build the ZIP→FIPS crosswalk (geometric)", 1)
P("Script: scripts/build_zcta.py", bold=True)
P("For every ZCTA in every state file, take a representative interior point of "
  "the ZIP polygon and run a point-in-polygon test against an STRtree spatial "
  "index of all county polygons. The first county that contains that point "
  "wins.")

CODE(
    "for each ZCTA polygon:\n"
    "    centroid = ZCTA.representative_point()\n"
    "    for i in county_tree.query(centroid):\n"
    "        if county_shapes[i].contains(centroid):\n"
    "            zip_to_fips[ZCTA] = county_fips[i]\n"
    "            break"
)

P("Outputs:")
BULLET("app/static/data/zcta/<state>.json — simplified ZIP polygons for the front-end map.")
BULLET("app/static/data/zip-by-county-2020.json — { fips: [zip5, …] } — the canonical ZIP→FIPS lookup, ~33k ZIPs.")

P("Caveat: ZIPs that physically straddle two counties are assigned to whichever "
  "county contains the ZIP centroid. One winner, no ties — a deliberate "
  "compromise that matches how the dashboard expects each ZIP to roll up.",
  italic=True)

# ── Stage 2: lift into Postgres ──────────────────────────────────────────
H("Stage 2 — Lift the crosswalk into Postgres", 1)
P("Script: scripts/backfill_foreclosure_county.py", bold=True)
P("Creates a permanent lookup table seeded from the JSON crosswalk plus Realtor "
  "names plus a static FIPS-prefix → state-abbreviation map.")

CODE(
    "CREATE TABLE IF NOT EXISTS zip_county_map (\n"
    "  postal_code  VARCHAR(5)  PRIMARY KEY,\n"
    "  county_fips  VARCHAR(5)  NOT NULL,\n"
    "  county_name  TEXT,\n"
    "  state_id     CHAR(2)\n"
    ");"
)

P("Population steps:")
BULLET("Read app/static/data/zip-by-county-2020.json → flatten to (zip, fips) pairs.")
BULLET("Pull (county_fips, county_name) from realtor_inventory_county at MAX(month_date_yyyymm).")
BULLET("Normalize names: 'lowndes, ga' → 'Lowndes', 'miami-dade, fl' → 'Miami-Dade'.")
BULLET("Derive state_id from the first two digits of the FIPS (hardcoded STATE_FIPS dict).")
BULLET("TRUNCATE zip_county_map and COPY all rows in.")

# ── Stage 3: backfill foreclosure_records ────────────────────────────────
H("Stage 3 — Enrich foreclosure_records", 1)
P("Same script. Only foreclosure_com rows with a missing county get filled — "
  "auction_com rows already have county names from the scraper.")

CODE(
    "UPDATE foreclosure_records fr\n"
    "SET county = zcm.county_name\n"
    "FROM zip_county_map zcm\n"
    "WHERE zcm.postal_code = LPAD(fr.zip, 5, '0')\n"
    "  AND fr.source = 'foreclosure_com'\n"
    "  AND (fr.county IS NULL OR fr.county = '')\n"
    "  AND zcm.county_name IS NOT NULL;"
)

# ── Join keys ─────────────────────────────────────────────────────────────
H("Join keys after the pipeline", 1)

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Table"
hdr[1].text = "ZIP key"
hdr[2].text = "County key"

rows = [
    ("foreclosure_records",       "zip (5-digit, LPAD when joining)", "county (name only — no FIPS column today)"),
    ("zip_county_map",            "postal_code (PK)",                 "county_fips (the bridge)"),
    ("realtor_inventory_zip",     "postal_code",                      "—"),
    ("realtor_hotness_zip",       "postal_code",                      "—"),
    ("realtor_inventory_county",  "—",                                "county_fips"),
    ("realtor_hotness_county",    "—",                                "county_fips"),
    ("realtor_inventory_state",   "—",                                "state_id (2-letter)"),
]
for r in rows:
    cells = table.add_row().cells
    cells[0].text = r[0]
    cells[1].text = r[1]
    cells[2].text = r[2]

# ── Example join ─────────────────────────────────────────────────────────
H("Example: foreclosures × Realtor county metrics", 1)
P("Because foreclosure_records does not yet have a county_fips column, county-"
  "level joins go through zip_county_map:")

CODE(
    "SELECT f.id, f.zip, f.classification,\n"
    "       rc.median_listing_price, rc.median_days_on_market,\n"
    "       hc.hotness_score\n"
    "FROM foreclosure_records f\n"
    "JOIN zip_county_map zcm\n"
    "  ON zcm.postal_code = LPAD(f.zip, 5, '0')\n"
    "JOIN realtor_inventory_county rc\n"
    "  ON rc.county_fips = zcm.county_fips\n"
    " AND rc.month_date_yyyymm = (\n"
    "       SELECT MAX(month_date_yyyymm) FROM realtor_inventory_county)\n"
    "LEFT JOIN realtor_hotness_county hc\n"
    "  ON hc.county_fips = zcm.county_fips\n"
    " AND hc.month_date_yyyymm = rc.month_date_yyyymm\n"
    "WHERE f.status = 'active';"
)

# ── Recommended next step ────────────────────────────────────────────────
H("Recommended next step (optional)", 1)
P("Add a county_fips VARCHAR(5) column on foreclosure_records and have the "
  "same backfill script populate it alongside county. That removes the "
  "zip_county_map hop on every county-level query and gives a clean direct "
  "join: realtor_inventory_county.county_fips = foreclosure_records.county_fips.")

CODE(
    "ALTER TABLE foreclosure_records\n"
    "  ADD COLUMN IF NOT EXISTS county_fips VARCHAR(5);\n"
    "CREATE INDEX IF NOT EXISTS foreclosure_records_county_fips_idx\n"
    "  ON foreclosure_records(county_fips);"
)

# ── File index ───────────────────────────────────────────────────────────
H("File index", 1)
BULLET("data/counties-fips.geojson — county polygons (master shapefile).")
BULLET("data/_zcta_cache/<state>.geojson — cached per-state ZCTA polygons (downloaded).")
BULLET("app/static/data/zip-by-county-2020.json — ZIP→FIPS crosswalk JSON.")
BULLET("app/static/data/zcta/<state>.json — simplified ZIP polygons for the map.")
BULLET("scripts/build_zcta.py — builds the crosswalk (point-in-polygon).")
BULLET("scripts/backfill_foreclosure_county.py — loads zip_county_map + backfills foreclosure_records.county.")
BULLET("scripts/import_realtor_county_state.py — loads realtor_inventory_county / _state / hotness_county.")
BULLET("scripts/import_realtor_zip.py — loads realtor_inventory_zip / hotness_zip.")
BULLET("scripts/import_listings.py — aggregates classified foreclosure rows per ZIP/FIPS into listings.json / listings-zip.json.")

doc.save(OUT)
print(f"wrote {OUT.relative_to(ROOT)}  ({OUT.stat().st_size/1024:.0f} KB)")